# from dotenv import load_dotenv
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.embeddings.openai import OpenAIEmbeddings
# from langchain.vectorstores import FAISS
# from langchain.chat_models import ChatOpenAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.callbacks import get_openai_callback
# import os
# import json
# from concurrent.futures import ThreadPoolExecutor
# import tempfile
# import fitz

# load_dotenv()
# openai_api_key = os.getenv("OPENAI_API_KEY")

# # Initialize components only if API key is available
# if openai_api_key:
#     text_embeddings = OpenAIEmbeddings()
#     llm = ChatOpenAI(model="gpt-4o")

# def extract_text(files):
#     for file in files:
#         if file.type == "application/pdf":
#             with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#                 tmp_file.write(file.getvalue())
#                 tmp_file_path = tmp_file.name

#             pdf_document = fitz.open(tmp_file_path)
#             all_text = ""
#             for page in pdf_document:
#                 all_text += page.get_text("text")
#             return all_text
#         if file.type == "application/msword" or file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             from docx import Document
#             doc = Document(file)
#             all_text = ""
#             for para in doc.paragraphs:
#                 all_text += para.text
#             return all_text
#         elif file.type == "text/plain":
#             return file.read().decode("utf-8")
#         else:
#             return ""
    
# def create_embeddings(text):
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000,   
#         chunk_overlap=200,  
#         length_function=len 
#     )
#     chunks = text_splitter.split_text(text=text)
#     vectorstore = FAISS.from_texts(chunks, embedding=text_embeddings)

#     return vectorstore

# def process_question(question, vectorstore, llm):
#     docs = vectorstore.similarity_search(query=question, k=3)
#     chain = load_qa_chain(llm=llm, chain_type="stuff")
#     with get_openai_callback() as cb:
#         response = chain.run(input_documents=docs, question=question)
#     return question, response

# def get_responses(files, prompts):
#     pdf_text = extract_text(files)
#     vectorstore = create_embeddings(pdf_text)

#     with ThreadPoolExecutor() as executor:
#         # List comprehension to submit tasks and gather responses
#         responses = executor.map(lambda q: process_question(q, vectorstore, llm), prompts)

#     question_response_map = {}
#     for question, response in responses:
#         question_response_map[question] = response


#     return question_response_map

#     # return response

# file_processing.py
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import ConversationalRetrievalChain
from langchain.callbacks import get_openai_callback
import os
import json
from concurrent.futures import ThreadPoolExecutor
import tempfile
import fitz

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Initialize components only if API key is available
if openai_api_key:
    text_embeddings = OpenAIEmbeddings()
    llm = ChatOpenAI(model="gpt-4o")

def extract_text(files):
    all_text = ""
    for file in files:
        if file.type == "application/pdf":
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(file.getvalue())
                tmp_file_path = tmp_file.name

            pdf_document = fitz.open(tmp_file_path)
            for page in pdf_document:
                all_text += page.get_text("text")

        elif file.type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            from docx import Document
            doc = Document(file)
            for para in doc.paragraphs:
                all_text += para.text + "\n"

        elif file.type == "text/plain":
            all_text += file.read().decode("utf-8") + "\n"

    return all_text

def create_embeddings(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text=text)
    vectorstore = FAISS.from_texts(chunks, embedding=text_embeddings)
    return vectorstore

def process_question(question, vectorstore, llm):
    docs = vectorstore.similarity_search(query=question, k=3)
    chain = load_qa_chain(llm=llm, chain_type="stuff")
    with get_openai_callback() as cb:
        response = chain.run(input_documents=docs, question=question)
    return question, response

def get_responses(files, prompts):
    pdf_text = extract_text(files)
    vectorstore = create_embeddings(pdf_text)

    with ThreadPoolExecutor() as executor:
        responses = executor.map(lambda q: process_question(q, vectorstore, llm), prompts)

    question_response_map = {}
    for question, response in responses:
        question_response_map[question] = response

    return question_response_map

def get_chatbot_chain(files):
    full_text = extract_text(files)
    vectorstore = create_embeddings(full_text)
    return ConversationalRetrievalChain.from_llm(llm=llm, retriever=vectorstore.as_retriever())
