import os
import io
import tempfile
import shutil
from typing import List, Dict, Any
import json
from uuid import uuid4
from docx import Document
import pdfplumber
from fastapi.middleware.cors import CORSMiddleware

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from pydantic import BaseModel
from openai import OpenAI as OpenAIClient
from dotenv import load_dotenv


from llama_parse import LlamaParse
from llama_index.llms.openai import OpenAI
from llama_index.core.node_parser import MarkdownElementNodeParser
from llama_index.core.memory import ChatMemoryBuffer
from llama_index.core import (
    load_index_from_storage,
    VectorStoreIndex,
    StorageContext,
)
from llama_index.core.memory import Memory
from llama_index.core.llms import ChatMessage
from llama_index.core import Settings
from llama_index.llms.openai import OpenAI as LlamaOpenAI
from llama_index.core.chat_engine import ContextChatEngine

load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("OPENAI_API_KEY is not set in .env file")

llm = LlamaOpenAI(model="gpt-4o", api_key=openai_api_key)
Settings.llm = llm
openai_client = OpenAIClient(api_key=openai_api_key)

llama_cloud_api_key = os.getenv("LLAMA_CLOUD_API_KEY")
if not llama_cloud_api_key:
    raise ValueError("LLAMA_CLOUD_API_KEY is not set in .env file")


def extract_text_from_folder(folder_path: str):
    file_paths = [
            os.path.join(folder_path, filename)
            for filename in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, filename))
        ]
    full_text = ""
    for path in file_paths:
        full_text += extract_text(path) + "\n\n"
    return full_text


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".doc":
        raise ValueError(".doc format not supported directly; use .docx instead.")
    else:
        raise ValueError("Unsupported file format: " + ext)

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text.strip()

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

async def create_embedding(uploaded_files, persist_dir="embeddings/runbooks_index"):
    all_nodes = []

    llama_parser = LlamaParse(result_type="markdown", api_key=llama_cloud_api_key)
    node_parser = MarkdownElementNodeParser(llm=llm, num_workers=15)

    for file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        documents = llama_parser.load_data(tmp_path)
        for doc in documents:
            doc.metadata["source"] = file.filename

        nodes = node_parser.get_nodes_from_documents(documents)
        base_nodes, objects = node_parser.get_nodes_and_objects(nodes)

        all_nodes.extend(base_nodes)
        all_nodes.extend(objects)

        os.unlink(tmp_path)

    print("Total nodes to index:", len(all_nodes))
    index = VectorStoreIndex(all_nodes)
    index.storage_context.persist(persist_dir)

    return persist_dir

def get_openai_chat_completion_repsonse(prompt: str, context_text: str):
    completion = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Runbook Context:\n\n{context_text}"}
            ],
            temperature=0.5,
        )
    response = completion.choices[0].message.content.strip()
    return response

def add_dr_steps_to_index(session_id: str, dr_steps: str, persist_dir="embeddings/runbooks_index"):
    doc = Document(
        text=dr_steps,
        metadata={"source": "DR_Steps", "session_id": session_id}
    )

    # Load storage context and index
    storage_context = StorageContext.from_defaults(persist_dir=persist_dir)
    index = load_index_from_storage(storage_context)

    # Insert as new node
    index.insert(doc)
    index.storage_context.persist(persist_dir)



def get_query_engine(embedding_path="embeddings/runbooks_index"):
    storage_context = StorageContext.from_defaults(persist_dir=embedding_path)
    vectorstore = load_index_from_storage(storage_context=storage_context)
    print("Docs in index:", len(vectorstore.docstore.docs))
    query_engine = vectorstore.as_query_engine(similarity_top_k=5)

    return query_engine

def get_chat_engine(memory, system_prompt="", embedding_path="embeddings/runbooks_index"):
    storage_context = StorageContext.from_defaults(persist_dir=embedding_path)
    vectorstore = load_index_from_storage(storage_context=storage_context)
    chat_engine = vectorstore.as_chat_engine(
            chat_mode="context",
            memory=memory,
            system_prompt=system_prompt,
        )
    return chat_engine

import re

def extract_json_from_response(response_text: str) -> str:
    """
    Extracts JSON content from a string wrapped in markdown-style triple backticks.
    """
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", response_text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return response_text.strip()  # fallback if no fences found
